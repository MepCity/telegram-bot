"""
Word ve Excel form doldurma modülü
"""
import os
import subprocess
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook
import config


class DocumentHandler:
    """Word ve Excel şablonlarını doldur"""
    
    def __init__(self):
        self.templates_dir = Path('gerek')
        self.template_dir = self.templates_dir  # Alias
        self.output_dir = Path('outputs')
        
    def fill_yetkilendirme_taahhutnamesi(self, tax_data, output_path=None):
        """
        Yetkilendirme Taahhütnamesi Word belgesini doldurur.
        
        Args:
            tax_data: PDF'den çıkarılan vergi bilgileri
            output_path: Çıktı dosyasının kaydedileceği yol (opsiyonel)
            
        Returns:
            str: Oluşturulan dosyanın yolu
        """
        try:
            # Şablon dosyasını aç - .docx formatında olmalı
            template_path = os.path.join(self.template_dir, 'YetkilendirmeTaahhutname.docx')
            
            # Eğer .docx yoksa .doc'u dönüştür
            if not os.path.exists(template_path):
                doc_path = os.path.join(self.template_dir, 'YetkilendirmeTaahhutname.doc')
                if os.path.exists(doc_path):
                    print("📄 .doc dosyasını .docx'e dönüştürüyorum...")
                    import subprocess
                    result = subprocess.run([
                        'soffice', '--headless', '--convert-to', 'docx',
                        '--outdir', self.template_dir, doc_path
                    ], capture_output=True, text=True)
                    if result.returncode != 0:
                        raise Exception(f"Dönüştürme hatası: {result.stderr}")
            
            doc = Document(template_path)
            
            # Güncel tarihi al (gün/ay/yıl formatında)
            current_date = datetime.now().strftime('%d/%m/%Y')
            
            # Belgedeki tarihleri güncelle
            # 1. Sağ üst köşedeki tarih (genellikle ilk paragraf)
            # 2. "TAAHHÜDÜN BAŞLANGIÇ TARİHİ" satırındaki tarih
            for paragraph in doc.paragraphs:
                text = paragraph.text
                # Tarih formatını ara ve değiştir (dd/mm/yyyy veya dd.mm.yyyy)
                import re
                # dd/mm/yyyy formatını bul
                if re.search(r'\d{2}/\d{2}/\d{4}', text):
                    # Sadece tarihi değiştir, diğer metni koru
                    new_text = re.sub(r'\d{2}/\d{2}/\d{4}', current_date, text)
                    paragraph.text = new_text
                    print(f"✓ Tarih güncellendi: {text.strip()} → {new_text.strip()}")
            
            # . . . işaretlerini bul ve değiştir (nokta boşluk nokta boşluk nokta)
            # İlk . . . → Vergi numarası
            # İkinci . . . → Firma unvanı
            replacements_made = 0
            
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                
                # Paragraftaki tüm . . . işaretlerini kontrol et
                while '. . .' in paragraph.text and replacements_made < 2:
                    # İlk . . . → vergi numarası
                    if replacements_made == 0 and tax_data.get('tax_number'):
                        paragraph.text = paragraph.text.replace('. . .', tax_data['tax_number'], 1)
                        replacements_made = 1
                        print(f"✓ Vergi numarası yerleştirildi: {tax_data['tax_number']}")
                    # İkinci . . . → firma unvanı
                    elif replacements_made == 1 and tax_data.get('company_name'):
                        paragraph.text = paragraph.text.replace('. . .', tax_data['company_name'], 1)
                        replacements_made = 2
                        print(f"✓ Firma unvanı yerleştirildi: {tax_data['company_name']}")
                    else:
                        break  # Veri yoksa dur
            
            # Sayfa düzenini ayarla (1 sayfaya sığdır)
            for section in doc.sections:
                # Margin'leri minimum seviyeye çek (yazıcı marjinleri)
                section.top_margin = Pt(22)      # ~0.75 cm (minimum)
                section.bottom_margin = Pt(22)   # ~0.75 cm (minimum)
                section.left_margin = Pt(36)     # ~1.25 cm
                section.right_margin = Pt(36)    # ~1.25 cm
            
            # Paragraf ve font ayarlarını optimize et
            for paragraph in doc.paragraphs:
                # Paragraf öncesi/sonrası boşlukları tamamen kaldır
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)  # Sıfır boşluk
                # Satır aralığını daha da azalt
                paragraph.paragraph_format.line_spacing = 0.9  # Maksimum sıkıştırma
                
                # Boş paragrafları tamamen kaldır (keep_together ile)
                if not paragraph.text.strip():
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 0.5  # Boş satırları minimize et
                
                # Font boyutunu küçült (eğer çok büyükse)
                for run in paragraph.runs:
                    if run.font.size and run.font.size > Pt(11):
                        run.font.size = Pt(10.5)  # Biraz daha küçült
                    elif run.font.size and run.font.size > Pt(10):
                        run.font.size = Pt(9.5)  # Daha kompakt
            
            # Çıktı dosyasını kaydet
            if output_path is None:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_path = os.path.join(self.output_dir, f'Yetkilendirme_Taahhutnamesi_{timestamp}.docx')
            
            doc.save(output_path)
            print(f"✅ Yetkilendirme Taahhütnamesi oluşturuldu: {output_path}")
            return output_path
            
        except Exception as e:
            print(f"❌ Yetkilendirme Taahhütnamesi oluşturulurken hata: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def fill_kullanici_yetkilendirme_formu(self, tax_data, email, output_path=None):
        """
        Kullanıcı Yetkilendirme Formu Excel belgesini doldurur.
        
        Args:
            tax_data: PDF'den çıkarılan vergi bilgileri
            email: Firma e-posta adresi
            output_path: Çıktı dosyasının kaydedileceği yol (opsiyonel)
            
        Returns:
            str: Oluşturulan dosyanın yolu
        """
        try:
            # Şablon dosyasını aç
            template_path = os.path.join(self.template_dir, 'Kullanıcı Yetkilendirme Formu.xlsx')
            wb = load_workbook(template_path)
            ws = wb.active
            
            # Excel yapısına göre direkt hücreleri doldur
            # E6: Firma Adı/Unvanı
            ws['E6'] = tax_data.get('company_name', '')
            print(f"✓ Firma unvanı E6'ya yazıldı: {tax_data.get('company_name', '')[:50]}")
            
            # E7: Vergi No
            ws['E7'] = tax_data.get('tax_number', '')
            print(f"✓ Vergi numarası E7'ye yazıldı: {tax_data.get('tax_number', '')}")
            
            # E8: Adres
            ws['E8'] = tax_data.get('address', '')
            print(f"✓ Adres E8'e yazıldı: {tax_data.get('address', '')[:50]}")
            
            # E9: E-posta Adresi
            ws['E9'] = email
            print(f"✓ E-posta E9'a yazıldı: {email}")
            
            # E11-E19: Hatice Arslan bilgileri (şablonda zaten dolu, değiştirmiyoruz)
            
            # Çıktı dosyasını kaydet
            if output_path is None:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_path = os.path.join(self.output_dir, f'Kullanici_Yetkilendirme_Formu_{timestamp}.xlsx')
            
            wb.save(output_path)
            print(f"✅ Kullanıcı Yetkilendirme Formu oluşturuldu: {output_path}")
            return output_path
            
        except Exception as e:
            print(f"❌ Kullanıcı Yetkilendirme Formu oluşturulurken hata: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def convert_to_pdf(self, input_path):
        """
        Word veya Excel dosyasını PDF'e çevir (LibreOffice kullanarak)
        
        Args:
            input_path (str): Dönüştürülecek dosya
            
        Returns:
            str: PDF dosyasının yolu veya None
        """
        input_path = Path(input_path)
        output_dir = input_path.parent
        
        try:
            # LibreOffice yolunu bul (Railway/Linux için)
            soffice_paths = [
                '/usr/bin/soffice',
                '/usr/bin/libreoffice',
                '/usr/lib/libreoffice/program/soffice',  # Ubuntu alternatif yol
                '/opt/libreoffice/program/soffice',  # Manuel kurulum
                '/Applications/LibreOffice.app/Contents/MacOS/soffice',
                'soffice',
            ]
            
            soffice_cmd = None
            print(f"🔍 LibreOffice aranıyor...")
            
            for path in soffice_paths:
                if Path(path).exists():
                    soffice_cmd = path
                    print(f"✅ LibreOffice bulundu: {path}")
                    break
                else:
                    print(f"❌ Bulunamadı: {path}")
            
            if not soffice_cmd:
                # which ile ara
                try:
                    result = subprocess.run(['which', 'soffice'], capture_output=True, text=True, timeout=5)
                    if result.returncode == 0 and result.stdout.strip():
                        soffice_cmd = result.stdout.strip()
                        print(f"✅ which ile bulundu: {soffice_cmd}")
                    else:
                        # whereis ile de dene
                        result = subprocess.run(['whereis', 'soffice'], capture_output=True, text=True, timeout=5)
                        print(f"🔍 whereis soffice: {result.stdout}")
                        
                        # dpkg ile paket kontrolü
                        result = subprocess.run(['dpkg', '-L', 'libreoffice-writer'], capture_output=True, text=True, timeout=5)
                        print(f"📦 libreoffice-writer dosyaları:\n{result.stdout[:500]}")
                except Exception as e:
                    print(f"⚠️ which/whereis hatası: {e}")
                
                if not soffice_cmd:
                    print('❌ LibreOffice (soffice) hiçbir yerde bulunamadı')
                    return None
            
            # LibreOffice ile PDF'e çevir
            print(f"🔄 PDF dönüştürme başlıyor: {soffice_cmd}")
            result = subprocess.run([
                soffice_cmd,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', str(output_dir),
                str(input_path)
            ], capture_output=True, timeout=30, text=True)
            
            if result.returncode == 0:
                pdf_path = input_path.with_suffix('.pdf')
                if pdf_path.exists():
                    print(f'✅ PDF oluşturuldu: {pdf_path}')
                    return str(pdf_path)
                else:
                    print(f'❌ PDF dosyası oluşmadı: {pdf_path}')
            else:
                print(f'❌ LibreOffice dönüştürme hatası (returncode={result.returncode})')
                print(f'STDOUT: {result.stdout}')
                print(f'STDERR: {result.stderr}')
            
            return None
            
        except Exception as e:
            print(f'❌ PDF conversion exception: {e}')
            import traceback
            traceback.print_exc()
            return None
    
    def fill_sozlesme(self, tax_data, proje_turu, ucret_bilgisi, output_path=None):
        """
        Sözleşme Word belgesini doldurur.
        
        Args:
            tax_data: PDF'den çıkarılan vergi bilgileri
            proje_turu: Proje türü (ör: "TÜBİTAK 1501 PROJESİ")
            ucret_bilgisi: Ücret bilgisi dict {
                'tutar': str,  # Örn: "80.000 TL'nin %5'i"
                'aciklama': str  # Örn: "projenin onaylanması ile onaylanan tutar üzerinden %5'i"
            }
            output_path: Çıktı dosyasının kaydedileceği yol (opsiyonel)
            
        Returns:
            str: Oluşturulan dosyanın yolu
        """
        try:
            # Şablon dosyasını aç - .docx formatında olmalı
            template_path = os.path.join(self.template_dir, 'Sözleşme.docx')
            
            if not os.path.exists(template_path):
                raise Exception(f"Şablon bulunamadı: {template_path}")
            
            doc = Document(template_path)
            
            # Firma adı ve vergi numarası
            company_name = tax_data.get('company_name', '')
            tax_number = tax_data.get('tax_number', '')
            address = tax_data.get('address', '')
            
            # Ücret tutarını TL ile birlikte formatla
            tutar_with_tl = ucret_bilgisi['tutar'] + ' TL'
            
            # Metni değiştir
            for para in doc.paragraphs:
                # 1. ... (....) → Firma adı (Vergi No)
                if '... (....)' in para.text:
                    para.text = para.text.replace('... (....)', f'{company_name} ({tax_number})')
                    print(f"✓ Firma bilgisi yerleştirildi: {company_name} ({tax_number})")
                
                # 2. ... hazırlanması → Proje türü
                if '... hazırlanması' in para.text:
                    para.text = para.text.replace('... hazırlanması', f'{proje_turu.upper()} hazırlanması')
                    print(f"✓ Proje türü yerleştirildi: {proje_turu.upper()}")
                
                # 3. Madde 3-1: ..... adresindeki → Firma adresi
                if '..... adresindeki' in para.text:
                    para.text = para.text.replace('..... adresindeki', f'{address} adresindeki')
                    print(f"✓ Adres yerleştirildi: {address}")
                
                # 4. ... TL projenin → Ücret tutarı (TL ile birlikte)
                if '... TL projenin' in para.text:
                    para.text = para.text.replace('... TL projenin', f"{tutar_with_tl} projenin")
                    print(f"✓ Ücret tutarı yerleştirildi: {tutar_with_tl}")
                
                # 5. ... talep edilir → Ücret açıklaması
                if '... talep edilir' in para.text:
                    para.text = para.text.replace('... talep edilir', f"{ucret_bilgisi['aciklama']} talep edilir")
                    print(f"✓ Ücret açıklaması yerleştirildi: {ucret_bilgisi['aciklama']}")
                
                # 6. Sabit tutarın ... TL'si → Sabit tutarın 100.000 TL'si
                if 'Sabit tutarın ... TL' in para.text:
                    para.text = para.text.replace('Sabit tutarın ... TL', f"Sabit tutarın {tutar_with_tl}")
                    print(f"✓ Sabit tutar yerleştirildi: {tutar_with_tl}")
            
            # Çıktı dosyasını kaydet
            if not output_path:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                safe_company = company_name[:30].replace('/', '_').replace('\\', '_')
                filename = f'Sozlesme_{safe_company}_{timestamp}.docx'
                output_path = os.path.join(self.output_dir, filename)
            
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            
            print(f'✅ Sözleşme oluşturuldu: {output_path}')
            return output_path
            
        except Exception as e:
            print(f'❌ Sözleşme doldurma hatası: {e}')
            import traceback
            traceback.print_exc()
            return None
    
    def fill_kosgeb_vekaletname(self, tax_data, output_path=None):
        """
        KOSGEB Vekaletname belgesini doldurur.
        
        Args:
            tax_data: PDF'den çıkarılan vergi bilgileri (company_name, tax_number, address, email)
            output_path: Çıktı dosyasının kaydedileceği yol (opsiyonel)
            
        Returns:
            str: Oluşturulan dosyanın yolu
        """
        try:
            # Şablon dosyasını aç
            template_path = os.path.join(self.template_dir, 'kosgeb_vekaletname.docx')
            
            if not os.path.exists(template_path):
                raise FileNotFoundError(f"Şablon bulunamadı: {template_path}")
            
            doc = Document(template_path)
            
            # 10 yıl sonraki tarihi hesapla
            from datetime import timedelta
            future_date = datetime.now() + timedelta(days=365*10)
            
            # Türkçe ay isimleri
            turkish_months = {
                1: 'ocak', 2: 'şubat', 3: 'mart', 4: 'nisan',
                5: 'mayıs', 6: 'haziran', 7: 'temmuz', 8: 'ağustos',
                9: 'eylül', 10: 'ekim', 11: 'kasım', 12: 'aralık'
            }
            
            # Türkçe gün isimleri
            turkish_days = {
                0: 'Pazartesi', 1: 'Salı', 2: 'Çarşamba', 3: 'Perşembe',
                4: 'Cuma', 5: 'Cumartesi', 6: 'Pazar'
            }
            
            # Sayıyı yazıya çevir (basit versiyon 1-31 için)
            ones = ['', 'bir', 'iki', 'üç', 'dört', 'beş', 'altı', 'yedi', 'sekiz', 'dokuz']
            tens = ['', 'on', 'yirmi', 'otuz']
            
            def number_to_turkish(n):
                if n < 10:
                    return ones[n]
                elif n < 40:
                    tens_digit = n // 10
                    ones_digit = n % 10
                    return tens[tens_digit] + ones[ones_digit]
                return str(n)
            
            # Yılı yazıya çevir
            year = future_date.year
            thousands = year // 1000
            hundreds = (year % 1000) // 100
            tens_part = (year % 100) // 10
            ones_part = year % 10
            
            year_text = ''
            if thousands == 2:
                year_text = 'ikibin'
            if hundreds > 0:
                year_text += ones[hundreds] + 'yüz'
            if tens_part > 0:
                year_text += tens[tens_part]
            if ones_part > 0:
                year_text += ones[ones_part]
            
            # Formatlanmış tarih oluştur
            day_num = future_date.day
            month_num = future_date.month
            day_name = turkish_days[future_date.weekday()]
            
            date_str = f"{future_date.strftime('%d.%m.%Y')} ({number_to_turkish(day_num)} {turkish_months[month_num]} {year_text})  {day_name}"
            
            print(f"✓ Gelecek tarih: {date_str}")
            
            # Vergi levhasından alınan bilgiler
            company_name = tax_data.get('company_name', '')
            tax_number = tax_data.get('tax_number', '')
            address = tax_data.get('address', '')
            email = tax_data.get('email', '')  # Bot'tan alınacak
            
            # Belgedeki tarihleri ve alanları doldur
            for paragraph in doc.paragraphs:
                text = paragraph.text
                
                # 1. Tarihleri güncelle (31.12.2030 formatı)
                import re
                if re.search(r'31\.12\.2030 \(otuzbir aralık ikibinotuz\)\s+\w+', text):
                    paragraph.text = re.sub(
                        r'31\.12\.2030 \(otuzbir aralık ikibinotuz\)\s+\w+',
                        date_str,
                        text
                    )
                    print(f"✓ Tarih güncellendi")
                
                # 2. İlk "… vergi numaralı" → Vergi numarası ekle
                if '… vergi numaralı' in text and tax_number:
                    paragraph.text = text.replace('… vergi numaralı', f'{tax_number} vergi numaralı', 1)
                    print(f"✓ Vergi numarası yerleştirildi: {tax_number}")
                
                # 3. "… şirket adına" → Şirket adı ekle
                if '… şirket adına' in text and company_name:
                    paragraph.text = paragraph.text.replace('… şirket adına', f'{company_name} şirket adına', 1)
                    print(f"✓ Şirket adı yerleştirildi: {company_name}")
                
                # 4. Tek başına "…" olan paragrafları doldur
                if text.strip() == '…':
                    # Bir önceki paragrafın içeriğine göre karar ver
                    # Bu basit implementasyon - şirket adını yerleştir
                    if company_name:
                        paragraph.text = company_name
                        print(f"✓ Şirket adı (VEKİL EDEN) yerleştirildi")
                
                # 5. "Vergi Numarası: …"
                if text.startswith('Vergi Numarası: …') and tax_number:
                    paragraph.text = f'Vergi Numarası: {tax_number}'
                    print(f"✓ Vergi numarası alanı dolduruldu")
                
                # 6. "Adresi: …"
                if text.startswith('Adresi: …') and address:
                    paragraph.text = f'Adresi: {address}'
                    print(f"✓ Adres alanı dolduruldu")
                
                # 7. "Elektronik Posta Adresi: …"
                if text.startswith('Elektronik Posta Adresi: …') and email:
                    paragraph.text = f'Elektronik Posta Adresi: {email}'
                    print(f"✓ E-posta alanı dolduruldu")
            
            # Çıktı dosyasını kaydet
            if not output_path:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                safe_company = company_name[:30].replace('/', '_').replace('\\', '_') if company_name else 'firma'
                filename = f'KOSGEB_Vekaletname_{safe_company}_{timestamp}.docx'
                output_path = os.path.join(self.output_dir, filename)
            
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            
            print(f'✅ KOSGEB Vekaletname oluşturuldu: {output_path}')
            return output_path
            
        except Exception as e:
            print(f'❌ KOSGEB Vekaletname doldurma hatası: {e}')
            import traceback
            traceback.print_exc()
            return None


if __name__ == '__main__':
    # Test
    handler = DocumentHandler()
    
    test_data = {
        'company_name': 'STİLLA OTOMOTİV TURİZM İNŞAAT TEKSTİL İTHALAT VE İHRACAT SANAYİ LİMİTED ŞİRKETİ',
        'tax_number': '1234567890',
        'address': 'VEYSEL KARANİ MAH. ESKİ GEMLİK YOLU CAD. NO: 236 F'
    }
    
    # Word
    word_path = handler.fill_yetkilendirme_taahhutnamesi(test_data)
    print(f'Word oluşturuldu: {word_path}')
    
    # Excel
    excel_path = handler.fill_kullanici_yetkilendirme_formu(test_data, 'test@example.com')
    print(f'Excel oluşturuldu: {excel_path}')
